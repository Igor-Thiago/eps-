from dataclasses import dataclass
from pathlib import Path

from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from core.models import Evidencia, RelatorioGerado
from .hash_service import HashService


class CertidaoError(Exception):
    pass


@dataclass(frozen=True)
class CertidaoResult:
    docx_path: Path
    hash_sha256: str
    generated_at_utc: str


_GOLD = RGBColor(0xC8, 0xA9, 0x51)
_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_GREY = RGBColor(0xF5, 0xF5, 0xF5)


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_border_bottom(cell, hex_color: str = 'C8A951', size: str = '12'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), hex_color)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


class CertidaoService:
    def generate(self, caso, analista, config=None) -> CertidaoResult:
        output_path = Path(caso.path_pasta) / f'certidao_{caso.id:04d}.docx'
        generated_at = timezone.now()

        try:
            doc = self._build_docx(caso, analista, config, generated_at)
            doc.save(str(output_path))
        except Exception as exc:
            raise CertidaoError(f'Nao foi possivel gerar a certidao: {exc}') from exc

        return CertidaoResult(
            docx_path=output_path,
            hash_sha256=HashService.hash_file(output_path),
            generated_at_utc=generated_at.isoformat(),
        )

    # ── Document builder ──────────────────────────────────────────────

    def _build_docx(self, caso, analista, config, generated_at):
        doc = Document()
        self._configure_page(doc)

        cabecalho = config.cabecalho if config and config.cabecalho else 'POLÍCIA CIVIL DO DISTRITO FEDERAL\nCORF — Coordenadoria de Repressão aos Crimes Cibernéticos'

        # Logo + cabeçalho
        if config and config.logo and hasattr(config.logo, 'path'):
            logo_path = Path(config.logo.path)
            if logo_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(logo_path), width=Cm(3.5))

        self._heading_block(doc, cabecalho)
        self._title_block(doc, 'CERTIDÃO DE COLETA DE EVIDÊNCIAS DIGITAIS')

        # Seções
        self._section_dados_caso(doc, caso, analista, generated_at)
        self._section_resumo(doc, caso)
        self._section_sites(doc, caso)
        self._section_videos(doc, caso)
        self._section_screenshots(doc, caso)
        self._section_forenses(doc, caso)
        self._section_downloads(doc, caso)
        self._section_relatorios(doc, caso)
        self._section_assinatura(doc, caso, analista, config, generated_at)

        return doc

    # ── Page setup ────────────────────────────────────────────────────

    def _configure_page(self, doc):
        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Header / Title ────────────────────────────────────────────────

    def _heading_block(self, doc, cabecalho: str):
        for line in cabecalho.split('\n'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = _BLACK
        p_format = doc.paragraphs[-1].paragraph_format
        p_format.space_after = Pt(4)

    def _title_block(self, doc, title: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _GOLD
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(14)
        # Linha decorativa dourada via border bottom
        self._add_gold_rule(doc)

    def _add_gold_rule(self, doc):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'C8A951')
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Section helpers ───────────────────────────────────────────────

    def _section_title(self, doc, text: str):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = _GOLD
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    def _kv_table(self, doc, rows: list[tuple]):
        table = doc.add_table(rows=len(rows), cols=2)
        table.style = 'Table Grid'
        for i, (key, value) in enumerate(rows):
            row = table.rows[i]
            kc = row.cells[0]
            vc = row.cells[1]
            kc.width = Cm(5)
            vc.width = Cm(12)
            _set_cell_bg(kc, 'F0ECD8')
            kp = kc.paragraphs[0]
            kr = kp.add_run(str(key))
            kr.bold = True
            kr.font.size = Pt(8)
            kr.font.color.rgb = _BLACK
            vp = vc.paragraphs[0]
            vr = vp.add_run(str(value) if value else '—')
            vr.font.size = Pt(8)
            vr.font.color.rgb = _BLACK
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def _evidence_table(self, doc, headers: list[str], rows: list[list]):
        if not rows:
            p = doc.add_paragraph()
            run = p.add_run('Nenhuma evidência deste tipo registrada.')
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            return

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'

        # Header row
        hrow = table.rows[0]
        for j, header in enumerate(headers):
            cell = hrow.cells[j]
            _set_cell_bg(cell, '1A1A1A')
            _set_cell_border_bottom(cell)
            p = cell.paragraphs[0]
            run = p.add_run(header.upper())
            run.bold = True
            run.font.size = Pt(7)
            run.font.color.rgb = _WHITE

        # Data rows
        for i, row_data in enumerate(rows):
            row = table.rows[i + 1]
            bg = 'FFFFFF' if i % 2 == 0 else 'F8F8F8'
            for j, value in enumerate(row_data):
                cell = row.cells[j]
                _set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                run = p.add_run(str(value) if value else '—')
                run.font.size = Pt(7)
                run.font.color.rgb = _BLACK

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Content sections ──────────────────────────────────────────────

    def _section_dados_caso(self, doc, caso, analista, generated_at):
        self._section_title(doc, '1. Dados do Caso e do Analista')
        self._kv_table(doc, [
            ('Nome do caso', caso.nome),
            ('Nº do processo', caso.numero_processo or '—'),
            ('Pasta do caso', caso.path_pasta),
            ('Analista responsável', analista.get_full_name() or analista.matricula),
            ('Matrícula', analista.matricula),
            ('Data de abertura', caso.created_at.strftime('%d/%m/%Y %H:%M UTC')),
            ('Certidão gerada em', generated_at.strftime('%d/%m/%Y %H:%M UTC')),
        ])

    def _section_resumo(self, doc, caso):
        self._section_title(doc, '2. Resumo de Evidências Coletadas')
        contagens = {
            tipo.label: caso.evidencias.filter(tipo=tipo.value).count()
            for tipo in Evidencia.Tipo
        }
        rows = [(tipo, str(qtd)) for tipo, qtd in contagens.items()]
        rows.append(('Total geral', str(caso.evidencias.count())))
        self._kv_table(doc, rows)

    def _section_sites(self, doc, caso):
        self._section_title(doc, '3. Sites Preservados')
        evidencias = caso.evidencias.filter(tipo=Evidencia.Tipo.SITE).order_by('created_at')
        rows = [
            [
                e.metadados_json.get('url', '—'),
                e.created_at.strftime('%d/%m/%Y %H:%M'),
                e.hash_sha256[:20] + '…',
            ]
            for e in evidencias
        ]
        self._evidence_table(doc, ['URL', 'Data/Hora', 'SHA-256 (trunc.)'], rows)

    def _section_videos(self, doc, caso):
        self._section_title(doc, '4. Gravações de Tela')
        evidencias = caso.evidencias.filter(tipo=Evidencia.Tipo.VIDEO).order_by('created_at')
        rows = [
            [
                Path(e.path_arquivo).name,
                e.created_at.strftime('%d/%m/%Y %H:%M'),
                e.hash_sha256[:20] + '…',
            ]
            for e in evidencias
        ]
        self._evidence_table(doc, ['Arquivo', 'Data/Hora', 'SHA-256 (trunc.)'], rows)

    def _section_screenshots(self, doc, caso):
        self._section_title(doc, '5. Capturas de Tela')
        evidencias = caso.evidencias.filter(tipo=Evidencia.Tipo.SCREENSHOT).order_by('created_at')
        rows = [
            [
                Path(e.path_arquivo).name,
                e.metadados_json.get('capture_mode', '—'),
                e.created_at.strftime('%d/%m/%Y %H:%M'),
                e.hash_sha256[:20] + '…',
            ]
            for e in evidencias
        ]
        self._evidence_table(doc, ['Arquivo', 'Modo', 'Data/Hora', 'SHA-256 (trunc.)'], rows)

    def _section_forenses(self, doc, caso):
        self._section_title(doc, '6. Cópias Forenses de Arquivos')
        evidencias = caso.evidencias.filter(tipo=Evidencia.Tipo.COPIA_FORENSE).order_by('created_at')
        rows = [
            [
                e.metadados_json.get('original_name', '—'),
                str(e.metadados_json.get('size_bytes', '—')) + ' bytes',
                e.created_at.strftime('%d/%m/%Y %H:%M'),
                e.hash_sha256[:20] + '…',
            ]
            for e in evidencias
        ]
        self._evidence_table(doc, ['Arquivo Original', 'Tamanho', 'Data/Hora', 'SHA-256 (trunc.)'], rows)

    def _section_downloads(self, doc, caso):
        self._section_title(doc, '7. Downloads Monitorados')
        evidencias = caso.evidencias.filter(tipo=Evidencia.Tipo.DOWNLOAD).order_by('created_at')
        rows = [
            [
                Path(e.path_arquivo).name,
                e.metadados_json.get('origin', '—'),
                e.created_at.strftime('%d/%m/%Y %H:%M'),
                e.hash_sha256[:20] + '…',
            ]
            for e in evidencias
        ]
        self._evidence_table(doc, ['Arquivo', 'Origem', 'Data/Hora', 'SHA-256 (trunc.)'], rows)

    def _section_relatorios(self, doc, caso):
        self._section_title(doc, '8. Relatórios Gerados')
        relatorios = caso.relatorios.all().order_by('created_at')
        rows = [
            [
                r.get_tipo_display(),
                Path(r.path_arquivo).name,
                r.created_at.strftime('%d/%m/%Y %H:%M'),
                r.hash_sha256[:20] + '…',
            ]
            for r in relatorios
        ]
        self._evidence_table(doc, ['Tipo', 'Arquivo', 'Data/Hora', 'SHA-256 (trunc.)'], rows)

    def _section_assinatura(self, doc, caso, analista, config, generated_at):
        self._add_gold_rule(doc)
        self._section_title(doc, '9. Declaração e Assinatura')

        nome = (config.assinatura_nome if config and config.assinatura_nome else analista.get_full_name() or analista.matricula)
        cargo = (config.assinatura_cargo if config and config.assinatura_cargo else 'Analista')
        orgao = (config.assinatura_orgao if config and config.assinatura_orgao else 'CORF/PCDF')

        p = doc.add_paragraph()
        run = p.add_run(
            f'Eu, {nome}, {cargo}, matrícula {analista.matricula}, declaro que as evidências '
            f'digitais registradas nesta certidão foram coletadas e preservadas conforme '
            f'os procedimentos técnicos estabelecidos pela {orgao}, utilizando o sistema '
            f'EDIFRAUDS, e que os hashes SHA-256 registrados garantem a integridade '
            f'dos artefatos coletados.'
        )
        run.font.size = Pt(9)
        run.font.color.rgb = _BLACK
        p.paragraph_format.space_after = Pt(28)

        # Linha de assinatura
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f'{nome}\n{cargo} — {orgao}\n{generated_at.strftime("%d/%m/%Y")}')
        run2.bold = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = _BLACK
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
